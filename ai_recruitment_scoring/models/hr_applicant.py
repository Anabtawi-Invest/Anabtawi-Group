# -*- coding: utf-8 -*-
import base64
import io
import json
import logging
import re

from odoo import models, fields, api, _
from odoo.exceptions import UserError

_logger = logging.getLogger(__name__)

# Language display names for prompt engineering
LANGUAGE_MAP = {
    'en': 'English',
    'ar': 'Arabic',
    'fr': 'French',
    'es': 'Spanish',
    'de': 'German',
    'pt': 'Portuguese',
    'zh': 'Chinese',
    'ja': 'Japanese',
    'ko': 'Korean',
    'tr': 'Turkish',
    'id': 'Indonesian',
    'hi': 'Hindi',
}


class HrApplicant(models.Model):
    _inherit = 'hr.applicant'

    # -----------------------
    # AI Score Fields
    # -----------------------
    ai_score = fields.Float(
        string='AI Score',
        readonly=True,
        copy=False,
        help='Overall AI matching score (0–100), weighted average of all dimensions.',
    )
    ai_score_ids = fields.One2many(
        'ai.applicant.score',
        'applicant_id',
        string='Dimension Scores',
        readonly=True,
        copy=False,
    )

    # -----------------------
    # AI Generated Content
    # -----------------------
    ai_summary = fields.Html(
        string='AI Summary',
        readonly=True,
        copy=False,
        sanitize=True,
        help='AI-generated candidate profile summary.',
    )
    ai_questions = fields.Html(
        string='Interview Questions',
        readonly=True,
        copy=False,
        sanitize=True,
        help='AI-generated categorized interview questions.',
    )
    ai_skill_gaps = fields.Text(
        string='Skill Gap Data',
        readonly=True,
        copy=False,
        help='JSON data of skill gap analysis for visualization.',
    )
    ai_strengths = fields.Text(
        string='Key Strengths',
        readonly=True,
        copy=False,
    )
    ai_weaknesses = fields.Text(
        string='Areas of Concern',
        readonly=True,
        copy=False,
    )

    # -----------------------
    # Metadata
    # -----------------------
    ai_analyzed_date = fields.Datetime(
        string='Last AI Analysis',
        readonly=True,
        copy=False,
    )
    ai_cv_text = fields.Text(
        string='Extracted CV Text',
        readonly=True,
        copy=False,
        help='Cached text extracted from the uploaded CV.',
    )
    ai_score_label = fields.Char(
        string='Score Label',
        compute='_compute_ai_score_label',
        store=False,
    )

    @api.depends('ai_score')
    def _compute_ai_score_label(self):
        for rec in self:
            score = rec.ai_score
            if score >= 85:
                rec.ai_score_label = _('Excellent')
            elif score >= 70:
                rec.ai_score_label = _('Strong')
            elif score >= 50:
                rec.ai_score_label = _('Moderate')
            elif score >= 30:
                rec.ai_score_label = _('Weak')
            elif score > 0:
                rec.ai_score_label = _('Poor')
            else:
                rec.ai_score_label = _('Not Scored')

    # ===================================================================
    #  PUBLIC ACTIONS (Button Methods)
    # ===================================================================

    def action_ai_analyze(self):
        """Main action: Extract CV text → call AI → populate all fields."""
        self.ensure_one()
        company = self.env.company

        if not company.ai_api_key:
            raise UserError(_(
                'AI API Key is not configured. '
                'Go to Settings → AI Recruitment to set it up.'
            ))

        # Step 1: Extract CV text from attachments
        cv_text = self._extract_cv_text()
        if not cv_text:
            raise UserError(_(
                'No CV text could be extracted. '
                'Please attach a PDF or DOCX resume to this applicant.'
            ))

        # Step 2: Get dimensions for this job
        dimensions = self._get_scoring_dimensions()

        # Step 3: Build prompt and call AI
        prompt = self._build_scoring_prompt(cv_text, dimensions)
        raw_response = self._call_ai_provider(prompt)

        # Step 4: Parse AI response and update fields
        self._parse_and_apply_ai_response(raw_response, cv_text, dimensions)

        # Step 5: Send notifications
        self._send_analysis_notification()

        return {
            'type': 'ir.actions.client',
            'tag': 'display_notification',
            'params': {
                'title': _('AI Analysis Complete'),
                'message': _('Score: %.0f/100 — %s') % (self.ai_score, self.ai_score_label),
                'type': 'success',
                'sticky': False,
            },
        }

    def action_ai_generate_questions(self):
        """Generate interview questions only (without re-scoring)."""
        self.ensure_one()
        company = self.env.company

        if not company.ai_api_key:
            raise UserError(_(
                'AI API Key is not configured. '
                'Go to Settings → AI Recruitment to set it up.'
            ))

        cv_text = self.ai_cv_text or self._extract_cv_text()
        if not cv_text:
            raise UserError(_(
                'No CV text available. Please run a full AI Analysis first, '
                'or attach a PDF/DOCX resume.'
            ))

        prompt = self._build_questions_prompt(cv_text)
        raw_response = self._call_ai_provider(prompt)
        self._parse_and_apply_questions(raw_response)

        return {
            'type': 'ir.actions.client',
            'tag': 'display_notification',
            'params': {
                'title': _('Questions Generated'),
                'message': _('Interview questions have been updated.'),
                'type': 'success',
                'sticky': False,
            },
        }

    def action_clear_ai_data(self):
        """Clear all AI-generated data for this applicant."""
        self.ensure_one()
        # Remove dimension scores
        self.ai_score_ids.unlink()
        self.write({
            'ai_score': 0,
            'ai_summary': False,
            'ai_questions': False,
            'ai_skill_gaps': False,
            'ai_strengths': False,
            'ai_weaknesses': False,
            'ai_analyzed_date': False,
            'ai_cv_text': False,
        })

    # ===================================================================
    #  SCORING DIMENSIONS
    # ===================================================================

    def _get_scoring_dimensions(self):
        """Get the scoring dimensions for this applicant's job position.

        Returns a list of dicts:
        [{'id': dim_id, 'code': 'skills', 'name': 'Skills Match',
          'description': '...', 'weight': 35, 'icon': '🔧'}, ...]
        """
        self.ensure_one()
        dimensions = []

        if self.job_id and self.job_id.ai_dimension_line_ids:
            for line in self.job_id.ai_dimension_line_ids:
                dimensions.append({
                    'id': line.dimension_id.id,
                    'code': line.dimension_id.code,
                    'name': line.dimension_id.name,
                    'description': line.dimension_id.description or '',
                    'weight': line.weight,
                    'icon': line.dimension_id.icon or '📋',
                })
        else:
            # Fallback: use all default dimensions
            default_dims = self.env['ai.scoring.dimension'].search([
                ('is_default', '=', True),
                ('active', '=', True),
            ])
            for dim in default_dims:
                dimensions.append({
                    'id': dim.id,
                    'code': dim.code,
                    'name': dim.name,
                    'description': dim.description or '',
                    'weight': dim.weight,
                    'icon': dim.icon or '📋',
                })

        # If still empty, use hardcoded fallback
        if not dimensions:
            dimensions = [
                {'id': False, 'code': 'skills', 'name': 'Skills Match',
                 'description': 'How well the candidate skills match the job requirements.',
                 'weight': 35, 'icon': '🔧'},
                {'id': False, 'code': 'experience', 'name': 'Experience',
                 'description': 'Relevance and depth of work experience.',
                 'weight': 30, 'icon': '📊'},
                {'id': False, 'code': 'education', 'name': 'Education',
                 'description': 'Education level and relevance.',
                 'weight': 15, 'icon': '🎓'},
                {'id': False, 'code': 'cultural_fit', 'name': 'Cultural Fit',
                 'description': 'Alignment with company culture and values.',
                 'weight': 20, 'icon': '🤝'},
            ]

        return dimensions

    # ===================================================================
    #  NOTIFICATIONS
    # ===================================================================

    def _send_analysis_notification(self):
        """Send Odoo internal notification and optional email after analysis."""
        self.ensure_one()
        company = self.env.company
        applicant_name = self.partner_name or self.name or _('Applicant')

        # --- Odoo Internal Notification (always, if enabled) ---
        if company.ai_notify_odoo:
            notification_body = _(
                '🤖 <strong>AI Analysis Complete</strong><br/>'
                'Applicant: <strong>%s</strong><br/>'
                'Score: <strong>%.0f/100</strong> — %s<br/>'
                'Job: %s'
            ) % (
                applicant_name,
                self.ai_score,
                self.ai_score_label,
                self.job_id.name if self.job_id else _('Not set'),
            )

            # Notify the current user
            self.env.user.notify_info(
                message=notification_body,
                title=_('AI Recruitment'),
                sticky=False,
            ) if hasattr(self.env.user, 'notify_info') else None

            # Post a message on the applicant's chatter
            self.message_post(
                body=notification_body,
                subject=_('AI Analysis Complete — Score: %.0f/100') % self.ai_score,
                message_type='notification',
                subtype_xmlid='mail.mt_note',
            )

            # Notify additional configured users via Odoo inbox
            notify_users = company.ai_notify_user_ids - self.env.user
            if notify_users:
                channel_or_partners = notify_users.mapped('partner_id')
                if channel_or_partners:
                    self.message_notify(
                        partner_ids=channel_or_partners.ids,
                        body=notification_body,
                        subject=_('AI Analysis: %s — Score: %.0f/100') % (
                            applicant_name, self.ai_score
                        ),
                    )

        # --- Optional Email Notification ---
        if company.ai_notify_email:
            self._send_analysis_email(applicant_name)

    def _send_analysis_email(self, applicant_name):
        """Send an email notification about the completed AI analysis."""
        self.ensure_one()
        company = self.env.company

        email_recipients = company.ai_notify_user_ids | self.env.user
        partner_ids = email_recipients.mapped('partner_id').ids

        if not partner_ids:
            return

        # Build dimension scores table for the email body
        score_rows = ''
        for score_line in self.ai_score_ids:
            score_rows += (
                f'<tr>'
                f'<td style="padding:6px 12px;">{score_line.dimension_icon} {score_line.dimension_id.name}</td>'
                f'<td style="padding:6px 12px;text-align:center;"><strong>{score_line.score:.0f}</strong>/100</td>'
                f'<td style="padding:6px 12px;text-align:center;">{score_line.weight}%</td>'
                f'</tr>'
            )

        email_body = _("""
            <div style="font-family: Arial, sans-serif; max-width: 600px;">
                <h2 style="color: #714B67;">🤖 AI Recruitment Analysis Complete</h2>
                <table style="width:100%%; border-collapse:collapse; margin:16px 0;">
                    <tr style="background:#f5f5f5;">
                        <td style="padding:8px 12px;"><strong>Applicant</strong></td>
                        <td style="padding:8px 12px;">%(applicant)s</td>
                    </tr>
                    <tr>
                        <td style="padding:8px 12px;"><strong>Job Position</strong></td>
                        <td style="padding:8px 12px;">%(job)s</td>
                    </tr>
                    <tr style="background:#f5f5f5;">
                        <td style="padding:8px 12px;"><strong>Overall Score</strong></td>
                        <td style="padding:8px 12px;"><strong style="font-size:18px;color:#714B67;">%(score).0f/100</strong> — %(label)s</td>
                    </tr>
                </table>
                %(dimension_table)s
                <p style="color:#666;font-size:13px;">
                    Open the applicant record in Odoo to view the full AI analysis, summary, and interview questions.
                </p>
            </div>
        """) % {
            'applicant': applicant_name,
            'job': self.job_id.name if self.job_id else _('Not set'),
            'score': self.ai_score,
            'label': self.ai_score_label,
            'dimension_table': (
                f'<h3 style="color:#714B67;">Dimension Scores</h3>'
                f'<table style="width:100%; border-collapse:collapse; border:1px solid #ddd;">'
                f'<thead><tr style="background:#714B67;color:#fff;">'
                f'<th style="padding:8px 12px;text-align:left;">Dimension</th>'
                f'<th style="padding:8px 12px;text-align:center;">Score</th>'
                f'<th style="padding:8px 12px;text-align:center;">Weight</th>'
                f'</tr></thead><tbody>{score_rows}</tbody></table>'
            ) if score_rows else '',
        }

        self.message_notify(
            partner_ids=partner_ids,
            body=email_body,
            subject=_('AI Recruitment: %s — Score: %.0f/100') % (
                applicant_name, self.ai_score
            ),
            mail_auto_delete=False,
            email_layout_xmlid='mail.mail_notification_light',
        )

    # ===================================================================
    #  CV TEXT EXTRACTION
    # ===================================================================

    def _extract_cv_text(self):
        """Extract text from the latest PDF or DOCX attachment."""
        self.ensure_one()
        attachments = self.env['ir.attachment'].search([
            ('res_model', '=', 'hr.applicant'),
            ('res_id', '=', self.id),
            ('mimetype', 'in', [
                'application/pdf',
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'application/msword',
            ]),
        ], order='create_date desc', limit=1)

        if not attachments:
            # Also try candidate-linked attachments
            if hasattr(self, 'candidate_id') and self.candidate_id:
                attachments = self.env['ir.attachment'].search([
                    ('res_model', '=', 'hr.candidate'),
                    ('res_id', '=', self.candidate_id.id),
                    ('mimetype', 'in', [
                        'application/pdf',
                        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        'application/msword',
                    ]),
                ], order='create_date desc', limit=1)

        if not attachments:
            return ''

        attachment = attachments[0]
        file_data = base64.b64decode(attachment.datas)

        if attachment.mimetype == 'application/pdf':
            text = self._extract_text_from_pdf(file_data)
        else:
            text = self._extract_text_from_docx(file_data)

        # Cache the extracted text
        if text:
            self.ai_cv_text = text

        return text

    @staticmethod
    def _extract_text_from_pdf(file_data):
        """Extract text from PDF binary data using PyMuPDF (fitz)."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise UserError(_(
                'PyMuPDF is required for PDF parsing. '
                'Install it with: pip install pymupdf'
            ))

        text_parts = []
        try:
            doc = fitz.open(stream=file_data, filetype='pdf')
            for page in doc:
                text_parts.append(page.get_text())
            doc.close()
        except Exception as e:
            _logger.error('Failed to extract PDF text: %s', e)
            raise UserError(_(
                'Failed to read the PDF file. '
                'Please ensure it is a valid, text-based PDF.'
            ))

        return '\n'.join(text_parts).strip()

    @staticmethod
    def _extract_text_from_docx(file_data):
        """Extract text from DOCX binary data using python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise UserError(_(
                'python-docx is required for DOCX parsing. '
                'Install it with: pip install python-docx'
            ))

        text_parts = []
        try:
            doc = Document(io.BytesIO(file_data))
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
        except Exception as e:
            _logger.error('Failed to extract DOCX text: %s', e)
            raise UserError(_(
                'Failed to read the DOCX file. '
                'Please ensure it is a valid Word document.'
            ))

        return '\n'.join(text_parts).strip()

    # ===================================================================
    #  AI PROVIDER ABSTRACTION
    # ===================================================================

    def _call_ai_provider(self, prompt):
        """Call the configured AI provider and return the raw text response."""
        company = self.env.company
        provider = company.ai_provider

        if provider == 'gemini':
            return self._call_gemini(prompt)
        elif provider == 'openai':
            return self._call_openai(prompt)
        else:
            raise UserError(_('Unsupported AI provider: %s') % provider)

    def _call_gemini(self, prompt):
        """Call Google Gemini API."""
        company = self.env.company

        try:
            import google.generativeai as genai
        except ImportError:
            raise UserError(_(
                'Google Generative AI library is required. '
                'Install it with: pip install google-generativeai'
            ))

        try:
            genai.configure(api_key=company.ai_api_key)
            model = genai.GenerativeModel(company.ai_model or 'gemini-2.0-flash')

            generation_config = genai.types.GenerationConfig(
                temperature=company.ai_temperature or 0.3,
                max_output_tokens=4096,
            )

            response = model.generate_content(
                prompt,
                generation_config=generation_config,
            )
            return response.text

        except Exception as e:
            _logger.error('Gemini API error: %s', e)
            raise UserError(_(
                'AI Analysis failed. Gemini API error:\n%s'
            ) % str(e))

    def _call_openai(self, prompt):
        """Call OpenAI API."""
        company = self.env.company

        try:
            from openai import OpenAI
        except ImportError:
            raise UserError(_(
                'OpenAI library is required. '
                'Install it with: pip install openai'
            ))

        try:
            client = OpenAI(api_key=company.ai_api_key)
            response = client.chat.completions.create(
                model=company.ai_model or 'gpt-4o-mini',
                messages=[
                    {
                        'role': 'system',
                        'content': (
                            'You are an expert HR recruitment analyst. '
                            'Always respond with valid JSON only, no markdown.'
                        ),
                    },
                    {'role': 'user', 'content': prompt},
                ],
                temperature=company.ai_temperature or 0.3,
                max_tokens=4096,
            )
            return response.choices[0].message.content

        except Exception as e:
            _logger.error('OpenAI API error: %s', e)
            raise UserError(_(
                'AI Analysis failed. OpenAI API error:\n%s'
            ) % str(e))

    # ===================================================================
    #  PROMPT ENGINEERING
    # ===================================================================

    def _get_job_context(self):
        """Build job context dict for prompt construction."""
        job = self.job_id
        if not job:
            return {}

        return {
            'title': job.name or '',
            'description': job.description or '',
            'required_skills': job.ai_required_skills or '',
            'preferred_skills': job.ai_preferred_skills or '',
            'min_experience': job.ai_experience_years or 0,
            'min_education': dict(job._fields['ai_education_level'].selection).get(
                job.ai_education_level, 'None'
            ) if job.ai_education_level else 'None',
            'custom_criteria': job.ai_scoring_criteria or '',
            'question_focus': job.ai_question_focus or '',
        }

    def _build_scoring_prompt(self, cv_text, dimensions):
        """Build the full scoring + summary + questions prompt with dynamic dimensions."""
        company = self.env.company
        lang = LANGUAGE_MAP.get(company.ai_language, 'English')
        max_questions = company.ai_max_questions or 10
        job_ctx = self._get_job_context()

        # Build dimensions section for the prompt
        dim_lines = []
        dim_json_keys = []
        for dim in dimensions:
            dim_lines.append(
                f"  - {dim['icon']} {dim['name']} (code: \"{dim['code']}\", weight: {dim['weight']}%): "
                f"{dim['description'] or 'Evaluate this aspect of the candidate.'}"
            )
            dim_json_keys.append(f'        "{dim["code"]}": <0-100>')

        dimensions_text = '\n'.join(dim_lines)
        scores_json_template = ',\n'.join(dim_json_keys)

        prompt = f"""You are an expert HR recruitment analyst. Analyze the following CV against the job requirements and provide a comprehensive assessment.

RESPOND ONLY WITH VALID JSON. No markdown, no code blocks, no explanations outside JSON.

=== JOB POSITION ===
Title: {job_ctx.get('title', 'Not specified')}
Description: {job_ctx.get('description', 'Not specified')}
Required Skills: {job_ctx.get('required_skills', 'Not specified')}
Preferred Skills: {job_ctx.get('preferred_skills', 'Not specified')}
Minimum Experience: {job_ctx.get('min_experience', 0)} years
Minimum Education: {job_ctx.get('min_education', 'None')}
Custom Criteria: {job_ctx.get('custom_criteria', 'None')}

=== SCORING DIMENSIONS ===
Score the candidate on each of these dimensions (0-100):
{dimensions_text}

=== CANDIDATE CV TEXT ===
{cv_text[:8000]}

=== INSTRUCTIONS ===
Respond in {lang} language.
Provide your analysis as a JSON object with EXACTLY this structure:

{{
    "scores": {{
{scores_json_template},
        "overall": <weighted average based on dimension weights above>
    }},
    "summary": "<HTML formatted summary of 3-5 paragraphs analyzing the candidate fit. Use <p>, <strong>, <ul>, <li> tags.>",
    "strengths": "<Comma-separated list of 3-5 key strengths>",
    "weaknesses": "<Comma-separated list of 2-4 areas of concern or gaps>",
    "skill_gaps": {{
        "matched": ["skill1", "skill2"],
        "missing": ["skill3", "skill4"],
        "extra": ["skill5"]
    }},
    "questions": [
        {{
            "category": "technical|behavioral|role_specific|situational",
            "difficulty": "easy|medium|hard",
            "question": "<The interview question>",
            "purpose": "<Brief reason why this question matters>"
        }}
    ]
}}

Generate exactly {max_questions} interview questions distributed across categories.
Focus areas for questions: {job_ctx.get('question_focus', 'General assessment')}
"""
        return prompt

    def _build_questions_prompt(self, cv_text):
        """Build a prompt for generating interview questions only."""
        company = self.env.company
        lang = LANGUAGE_MAP.get(company.ai_language, 'English')
        max_questions = company.ai_max_questions or 10
        job_ctx = self._get_job_context()

        prompt = f"""You are an expert HR interviewer. Generate targeted interview questions for this candidate.

RESPOND ONLY WITH VALID JSON. No markdown, no code blocks.

=== JOB ===
Title: {job_ctx.get('title', 'Not specified')}
Required Skills: {job_ctx.get('required_skills', 'Not specified')}
Focus Areas: {job_ctx.get('question_focus', 'General assessment')}

=== CANDIDATE CV ===
{cv_text[:6000]}

=== INSTRUCTIONS ===
Respond in {lang}.
Generate exactly {max_questions} questions as a JSON array:

[
    {{
        "category": "technical|behavioral|role_specific|situational",
        "difficulty": "easy|medium|hard",
        "question": "<question text>",
        "purpose": "<why this question matters>"
    }}
]

Distribute questions: ~40% technical, ~25% behavioral, ~20% role-specific, ~15% situational.
"""
        return prompt

    # ===================================================================
    #  RESPONSE PARSING & FIELD POPULATION
    # ===================================================================

    def _parse_and_apply_ai_response(self, raw_response, cv_text, dimensions):
        """Parse AI JSON response and write to record fields using dynamic dimensions."""
        self.ensure_one()
        data = self._safe_parse_json(raw_response)

        scores = data.get('scores', {})
        questions = data.get('questions', [])
        skill_gaps = data.get('skill_gaps', {})

        # --- Save individual dimension scores ---
        # Remove old scores first
        self.ai_score_ids.unlink()

        score_vals_list = []
        total_weighted = 0
        total_weight = 0

        for dim in dimensions:
            dim_score = min(max(float(scores.get(dim['code'], 0)), 0), 100)
            dim_weight = dim['weight']
            total_weighted += dim_score * dim_weight
            total_weight += dim_weight

            score_vals = {
                'applicant_id': self.id,
                'score': dim_score,
                'weight': dim_weight,
            }
            # Link to dimension record if it exists
            if dim['id']:
                score_vals['dimension_id'] = dim['id']

            score_vals_list.append(score_vals)

        # Create dimension score records
        if score_vals_list:
            self.env['ai.applicant.score'].create(score_vals_list)

        # Calculate overall weighted score
        overall = (total_weighted / total_weight) if total_weight > 0 else 0
        # Also consider the AI's own overall if provided
        ai_overall = float(scores.get('overall', 0))
        final_overall = overall if overall > 0 else ai_overall

        # Build questions HTML
        questions_html = self._format_questions_html(questions)

        # Build summary — use AI's HTML or wrap plain text
        summary = data.get('summary', '')
        if summary and not summary.strip().startswith('<'):
            summary = f'<p>{summary}</p>'

        self.write({
            'ai_score': min(max(final_overall, 0), 100),
            'ai_summary': summary,
            'ai_questions': questions_html,
            'ai_skill_gaps': json.dumps(skill_gaps) if skill_gaps else False,
            'ai_strengths': data.get('strengths', ''),
            'ai_weaknesses': data.get('weaknesses', ''),
            'ai_analyzed_date': fields.Datetime.now(),
            'ai_cv_text': cv_text,
        })

    def _parse_and_apply_questions(self, raw_response):
        """Parse questions-only response and update the questions field."""
        self.ensure_one()
        data = self._safe_parse_json(raw_response)

        # Response could be a list directly or a dict with 'questions' key
        if isinstance(data, list):
            questions = data
        else:
            questions = data.get('questions', [])

        questions_html = self._format_questions_html(questions)
        self.write({
            'ai_questions': questions_html,
        })

    @staticmethod
    def _safe_parse_json(raw_text):
        """Safely parse JSON from AI response, handling markdown code blocks."""
        if not raw_text:
            return {}

        text = raw_text.strip()

        # Strip markdown code fences if present
        if text.startswith('```'):
            text = re.sub(r'^```(?:json)?\s*\n?', '', text)
            text = re.sub(r'\n?```\s*$', '', text)

        try:
            return json.loads(text)
        except json.JSONDecodeError:
            # Try to find JSON object or array in the text
            match = re.search(r'(\{[\s\S]*\}|\[[\s\S]*\])', text)
            if match:
                try:
                    return json.loads(match.group(1))
                except json.JSONDecodeError:
                    pass
            _logger.warning('Could not parse AI response as JSON: %s', text[:500])
            return {}

    @staticmethod
    def _format_questions_html(questions):
        """Format a list of question dicts into styled HTML."""
        if not questions:
            return '<p><em>No questions generated.</em></p>'

        category_icons = {
            'technical': '🔧',
            'behavioral': '🧠',
            'role_specific': '🎯',
            'situational': '💡',
        }
        category_labels = {
            'technical': 'Technical',
            'behavioral': 'Behavioral',
            'role_specific': 'Role-Specific',
            'situational': 'Situational',
        }
        difficulty_colors = {
            'easy': '#28a745',
            'medium': '#ffc107',
            'hard': '#dc3545',
        }

        # Group questions by category
        grouped = {}
        for q in questions:
            cat = q.get('category', 'technical')
            grouped.setdefault(cat, []).append(q)

        html_parts = []
        for cat_key in ['technical', 'behavioral', 'role_specific', 'situational']:
            cat_questions = grouped.get(cat_key, [])
            if not cat_questions:
                continue

            icon = category_icons.get(cat_key, '📋')
            label = category_labels.get(cat_key, cat_key.replace('_', ' ').title())

            html_parts.append(
                f'<h4 style="margin-top:16px;margin-bottom:8px;color:#495057;">'
                f'{icon} {label} Questions</h4>'
            )
            html_parts.append('<ol style="padding-left:20px;">')

            for q in cat_questions:
                diff = q.get('difficulty', 'medium')
                diff_color = difficulty_colors.get(diff, '#6c757d')
                purpose = q.get('purpose', '')
                question_text = q.get('question', '')

                html_parts.append(
                    f'<li style="margin-bottom:12px;">'
                    f'<strong>{question_text}</strong>'
                    f' <span style="background-color:{diff_color};color:#fff;'
                    f'padding:2px 8px;border-radius:12px;font-size:11px;'
                    f'margin-left:6px;">{diff.upper()}</span>'
                )
                if purpose:
                    html_parts.append(
                        f'<br/><em style="color:#6c757d;font-size:13px;">'
                        f'Purpose: {purpose}</em>'
                    )
                html_parts.append('</li>')

            html_parts.append('</ol>')

        return '\n'.join(html_parts)
